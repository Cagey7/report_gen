import re
import os
import psycopg2
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.shared import RGBColor
from utils.utils import format_month_range
from db.fetcher import TradeDataFetcher
from report_data.preparer import TradeDataPreparer


HEADER_BG = "1F4E79"  # темно-синий
HEADER_TEXT = RGBColor(255, 255, 255)  # белый
HEADER_BASIC = "D9D9D9"  # светло-серый
ROW_BG1 = "DCE6F1"    # светло-синий
ROW_BG2 = "FFFFFF"    # белый


class TradeDocumentGenerator:
    def __init__(self, prepared_data):
        self.prepared_data = prepared_data

    def generate(self):
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(2 / 2.54)
        section.bottom_margin = Inches(2 / 2.54)
        section.left_margin = Inches(3 / 2.54)
        section.right_margin = Inches(1.5 / 2.54)

        change_color = self.prepared_data["change_color"]

        args = [doc, self.prepared_data["document_header"]]
        if "category_description" in self.prepared_data:
            args.append(self.prepared_data["category_description"])

        self.add_document_header(*args)
        self.add_summary_paragraph(doc, self.prepared_data["summary_text"])
        self.add_summary_table(doc, self.prepared_data["summary_header"], self.prepared_data["summary_table"], change_color)

        if self.prepared_data.get("trade_dynamics_table"):
            self.generate_trade_dynamics_table(doc, self.prepared_data["trade_dynamics_table"], change_color)

        if self.prepared_data.get("months_table_data"):
            self.generate_trade_dynamics_table(doc, self.prepared_data["months_table_data"], change_color, self.prepared_data["end_year"])

        if self.prepared_data.get("country_table_data"):
            self.add_country_table(
                doc,
                self.prepared_data["country_table_data"],
                self.prepared_data["country_table_header"],
                self.prepared_data["country_table_units"],
                "country",
                change_color
            )

        if self.prepared_data.get("region_table_data"):
            self.add_country_table(
                doc,
                self.prepared_data["region_table_data"],
                self.prepared_data["region_table_header"],
                self.prepared_data["region_table_units"],
                "region",
                change_color
            ) 

        if self.prepared_data["summary_table"][1][2] == "0,0":
            return "Данных нет", "Данных нет", "Данных нет"
        
        self.add_import_analysis_text(doc, self.prepared_data["export_text"])
        self.add_import_analysis_text(doc, self.prepared_data["import_text"])

        self.generate_export_import_table(
            doc,
            self.prepared_data["export_header"],
            self.prepared_data["export_table"],
            self.prepared_data["months"],
            self.prepared_data["start_year"],
            self.prepared_data["end_year"],
            self.prepared_data["export_table_measure"],
            change_color
        )
        self.generate_export_import_table(
            doc,
            self.prepared_data["import_header"],
            self.prepared_data["import_table"],
            self.prepared_data["months"],
            self.prepared_data["start_year"],
            self.prepared_data["end_year"],
            self.prepared_data["import_table_measure"],
            change_color
        )

        return doc, self.prepared_data["filename"], self.prepared_data["short_filename"]

    def set_run_style(self, run):
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(14)

    def format_paragraph(self, paragraph, first_line_indent=True):
        if first_line_indent:
            paragraph.paragraph_format.first_line_indent = Inches(0.5)
        paragraph.paragraph_format.line_spacing = 1

    def add_document_header(self, doc, header_text, category_description=None):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = Cm(2)
        paragraph_format.right_indent = Cm(2)

        self.format_paragraph(paragraph, first_line_indent=False)

        run = paragraph.add_run(header_text)
        run.bold = True
        self.set_run_style(run)

        if category_description:
            run2 = paragraph.add_run("\n" + category_description)
            run2.italic = True
            self.set_run_style(run2)

    def add_summary_paragraph(self, doc, text):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.format_paragraph(paragraph)

        words = text.split(' ', 1)
        if not words:
            return

        run_first = paragraph.add_run(words[0] + (' ' if len(words) > 1 else ''))
        self.set_run_style(run_first)
        run_first.bold = True

        rest_text = words[1] if len(words) > 1 else ''

        start_idx = rest_text.find("составил")
        end_idx = rest_text.find("США")

        if start_idx != -1 and end_idx != -1 and end_idx > start_idx:
            before = rest_text[:start_idx + len("составил")]
            bold_part = rest_text[start_idx + len("составил"):end_idx + len("США")]
            after = rest_text[end_idx + len("США"):]

            if before:
                run_before = paragraph.add_run(before)
                self.set_run_style(run_before)

            if bold_part:
                run_bold = paragraph.add_run(bold_part)
                self.set_run_style(run_bold)
                run_bold.bold = True

            if after:
                run_after = paragraph.add_run(after)
                self.set_run_style(run_after)
        else:
            run_rest = paragraph.add_run(rest_text)
            self.set_run_style(run_rest)

    def set_cell_background(self, cell, color_hex):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color_hex)
        tcPr.append(shd)

    def set_cell_text_color(self, cell, rgb_color):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = rgb_color

    def add_cell_vertical_padding(self, cell, space_pts=4):
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = Pt(space_pts)
            paragraph.paragraph_format.space_after = Pt(space_pts)

    def get_available_page_width(self, doc):
        section = doc.sections[0]
        return section.page_width - section.left_margin - section.right_margin

    def add_summary_table(self, doc, title, table_data, change_color):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        self.format_paragraph(paragraph, first_line_indent=False)
        run = paragraph.add_run(title)
        run.italic = True
        self.set_run_style(run)

        rows = len(table_data)
        cols = len(table_data[0])
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, row in enumerate(table_data):
            for j, cell_text in enumerate(row):
                cell = table.cell(i, j)
                cell.text = str(cell_text)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
                    if i == 0 and j != 0:
                        run.bold = True
                    elif i == 1:
                        run.bold = True

                    if j == 3 and '-' in cell_text:
                        run.font.color.rgb = RGBColor(255, 0, 0)
                    if j == 3 and "ухудшился" in cell_text:
                        run.italic = True

                self.add_cell_vertical_padding(cell, space_pts=2)

                # change color
                if change_color:
                    if i == 0:
                        self.set_cell_background(cell, HEADER_BG)
                        self.set_cell_text_color(cell, HEADER_TEXT)
                    else:
                        self.set_cell_background(cell, ROW_BG1 if i % 2 == 1 else ROW_BG2)
                else:
                    if i == 0:
                        self.set_cell_background(cell, HEADER_BASIC)


    def add_import_analysis_text(self, doc, text_blocks):
        if text_blocks == "":
            return
        
        doc.add_paragraph().paragraph_format.space_after = Pt(10)

        for i, text in enumerate(text_blocks):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Inches(0.5)
            paragraph.paragraph_format.line_spacing = 1
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            if i == 0:
                first_word_match = re.match(r'\S+', text)
                first_word = first_word_match.group(0) if first_word_match else ''
                rest_text = text[len(first_word):]

                last_number_match = list(re.finditer(r'\d[\d\s.,]*', text))
                if last_number_match:
                    last_number = last_number_match[-1]
                    split_index = last_number.start()
                    middle_text = text[len(first_word):split_index]
                    end_text = text[split_index:]
                else:
                    middle_text = rest_text
                    end_text = ''

                run_first = paragraph.add_run(first_word)
                self.set_run_style(run_first)
                run_first.bold = True

                run_middle = paragraph.add_run(middle_text)
                self.set_run_style(run_middle)

                run_end = paragraph.add_run(end_text)
                self.set_run_style(run_end)
                run_end.bold = True

            elif i == 3 and ':' in text:
                before_colon, after_colon = text.split(':', 1)

                run_bold = paragraph.add_run(before_colon + ':')
                self.set_run_style(run_bold)
                run_bold.bold = True

                run_rest = paragraph.add_run(after_colon)
                self.set_run_style(run_rest)

            else:
                run = paragraph.add_run(text)
                self.set_run_style(run)
                if i in [1, 2]:
                    run.italic = True


    def generate_export_import_table(self, doc, table_header, table_data, months, start_year, end_year, units, change_color):
        if len(table_data) == 1:
            return
        for row in table_data:
            if len(row) > 9:
                del row[9]

        def merge_cells_vertically(table, col_idx, row_start, row_end):
            start_cell = table.cell(row_start, col_idx)
            for row in range(row_start + 1, row_end + 1):
                start_cell.merge(table.cell(row, col_idx))

        def merge_cells_horizontally(table, row_idx, col_start, col_end):
            start_cell = table.cell(row_idx, col_start)
            for col in range(col_start + 1, col_end + 1):
                start_cell.merge(table.cell(row_idx, col))

        def set_cell_background(cell, color):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), color)

            tcPr.append(shd)

        def set_cell_text_color(cell, rgb_color):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = rgb_color

        def center_cell(cell):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(3)
                paragraph.paragraph_format.space_after = Pt(3)
        
        def make_bold(cell):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        def set_repeat_table_header(row):
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            tblHeader = OxmlElement('w:tblHeader')
            trPr.append(tblHeader)

        doc.add_paragraph().paragraph_format.space_after = 0

        header_paragraph = doc.add_paragraph(table_header)
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_paragraph.runs[0]
        run.italic = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        header_paragraph.paragraph_format.space_after = 0

        table = doc.add_table(rows=2, cols=9)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row in table.rows:
            row.cells[0].width = Inches(2.5)

        set_repeat_table_header(table.rows[0])
        set_repeat_table_header(table.rows[1])

        cell = table.cell(0, 0)
        cell.text = "Товары"
        merge_cells_vertically(table, 0, 0, 1)
        if change_color:
            set_cell_background(cell, HEADER_BG)
            set_cell_text_color(cell, HEADER_TEXT)
        else:
            set_cell_background(cell, HEADER_BASIC)

        center_cell(cell)
        make_bold(cell)

        cell = table.cell(0, 1)
        cell.text = f"{start_year} год" if months[-1] == 12 else f"{format_month_range(months)}\n{start_year} {'год' if months[-1] == 12 else 'года'}"
        merge_cells_horizontally(table, 0, 1, 3)
        for i in range(1, 4):
            if change_color:
                set_cell_background(table.cell(0, i), HEADER_BG)
                set_cell_text_color(table.cell(0, i), HEADER_TEXT)
            else:
                set_cell_background(table.cell(0, i), HEADER_BASIC)

            center_cell(table.cell(0, i))
            make_bold(table.cell(0, i))

        cell = table.cell(0, 4)
        cell.text = f"{end_year} год" if months[-1] == 12 else f"{format_month_range(months)}\n{end_year} {'год' if months[-1] == 12 else 'года'}"
        merge_cells_horizontally(table, 0, 4, 6)
        for i in range(4, 7):
            if change_color:
                set_cell_background(table.cell(0, i), HEADER_BG)
                set_cell_text_color(table.cell(0, i), HEADER_TEXT)
            else:
                set_cell_background(table.cell(0, i), HEADER_BASIC)

            center_cell(table.cell(0, i))
            make_bold(table.cell(0, i))

        cell = table.cell(0, 7)
        cell.text = f"Прирост\n{start_year}/{end_year}"
        merge_cells_horizontally(table, 0, 7, 8)
        for i in range(7, 9):
            if change_color:
                set_cell_background(table.cell(0, i), HEADER_BG)
                set_cell_text_color(table.cell(0, i), HEADER_TEXT)
            else:
                set_cell_background(table.cell(0, i), HEADER_BASIC)

            center_cell(table.cell(0, i))
            make_bold(table.cell(0, i))

        headers = ["физ. объем", f"{units}$", "Доля", "физ. объем", f"{units}$", "Доля", "физ. объем", f"{units}$"]
        for i in range(8):
            cell = table.cell(1, i + 1)
            cell.text = headers[i]

            if change_color:
                set_cell_background(cell, HEADER_BG)
                set_cell_text_color(cell, HEADER_TEXT)
            else:
                set_cell_background(cell, HEADER_BASIC)

            center_cell(cell)
            make_bold(cell)

        # Данные
        for i, row_data in enumerate(table_data):
            row = table.add_row().cells
            for col, val in enumerate(row_data):
                cell = row[col]
                cell.text = str(val)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_before = Pt(3)
                    paragraph.paragraph_format.space_after = Pt(3)
                    run = paragraph.runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(9)

                    if i == 0:
                        run.bold = True
                    else:
                        if col in [2, 5]:
                            run.bold = True
                        if col in [1, 4]:
                            run.italic = True

                    if col in [7, 8]:
                        if val.strip().lower() == "new":
                            run.font.color.rgb = RGBColor(0, 176, 80)
                        elif val.strip().startswith("-"):
                            run.font.color.rgb = RGBColor(255, 0, 0)

                if change_color:
                    self.set_cell_background(cell, ROW_BG1 if (i + 2) % 2 == 1 else ROW_BG2)


    def add_country_table(self, doc, data, header, units, table_type, change_color):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(12)
        self.format_paragraph(paragraph, first_line_indent=False)
        run = paragraph.add_run(header)
        run.italic = True
        self.set_run_style(run)

        table = doc.add_table(rows=2 + len(data), cols=8)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        header_row_1 = table.rows[0].cells
        header_row_1[0].text = '№'
        if table_type == "region":
            header_row_1[1].text = 'Регионы'
        elif table_type == "country":
            header_row_1[1].text = 'Страны'
        header_row_1[2].text = f'Стоимость, {units}'
        header_row_1[5].text = 'Доля'

        header_row_2 = table.rows[1].cells
        header_row_2[2].text = 'ТО'
        header_row_2[3].text = 'Экспорт'
        header_row_2[4].text = 'Импорт'
        header_row_2[5].text = 'ТО'
        header_row_2[6].text = 'Экспорт'
        header_row_2[7].text = 'Импорт'

        def merge_horizontally(row, start_idx, end_idx):
            start_cell = row.cells[start_idx]
            for idx in range(start_idx + 1, end_idx + 1):
                start_cell.merge(row.cells[idx])

        def merge_vertically(table, col_idx, start_row_idx, end_row_idx):
            start_cell = table.cell(start_row_idx, col_idx)
            for row_idx in range(start_row_idx + 1, end_row_idx + 1):
                start_cell.merge(table.cell(row_idx, col_idx))

        merge_vertically(table, 0, 0, 1)
        merge_vertically(table, 1, 0, 1)
        merge_horizontally(table.rows[0], 2, 4)
        merge_horizontally(table.rows[0], 5, 7)

        # Заливка заголовков
        for row in [table.rows[0], table.rows[1]]:
            for cell in row.cells:
                if change_color:
                    self.set_cell_background(cell, HEADER_BG)
                    self.set_cell_text_color(cell, HEADER_TEXT)
                else:
                    self.set_cell_background(cell, HEADER_BASIC)

                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.runs[0]
                    run.font.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    paragraph.paragraph_format.space_before = Pt(2)
                    paragraph.paragraph_format.space_after = Pt(2)

        for row in table.rows:
            row.cells[0].width = Cm(1)
            row.cells[1].width = Cm(3)

        for row_idx, row_data in enumerate(data, start=2):
            row_cells = table.rows[row_idx].cells
            for col_idx, value in enumerate(row_data):
                cell = row_cells[col_idx]
                cell.text = str(value)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    if row_idx == 2:
                        run.bold = True
                    paragraph.paragraph_format.space_before = Pt(2)
                    paragraph.paragraph_format.space_after = Pt(2)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                # Чередование цвета строк
                if change_color:
                    self.set_cell_background(cell, ROW_BG1 if row_idx % 2 == 1 else ROW_BG2)

        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_before = Pt(2)
                    paragraph.paragraph_format.space_after = Pt(2)


    def generate_trade_dynamics_table(self, doc, table_data, change_color, year=None):
        if not table_data or len(table_data) < 2:
            return

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        self.format_paragraph(paragraph, first_line_indent=False)
        
        if year:
            title = f"Динамика товарооборота по месяцам за {year} год"
            font_size = Pt(10)
            if len(table_data[0]) < 8:
                font_size = Pt(12)
        else:
            title = "Динамика товарооборота по годам"
            font_size = Pt(12)
        run = paragraph.add_run(title)
        run.italic = True
        self.set_run_style(run)

        rows = len(table_data)
        cols = len(table_data[0])
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, row in enumerate(table_data):
            for j, value in enumerate(row):
                cell = table.cell(i, j)
                cell.text = str(value)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = font_size
                    if i == 0:
                        run.bold = True
                    elif i == 1:
                        run.bold = True

                # change color
                if change_color:
                    if i == 0:
                        self.set_cell_background(cell, HEADER_BG)
                        self.set_cell_text_color(cell, HEADER_TEXT)
                    else:
                        self.set_cell_background(cell, ROW_BG1 if i % 2 == 1 else ROW_BG2)
                else:
                    if i == 0:
                        self.set_cell_background(cell, HEADER_BASIC)

                self.add_cell_vertical_padding(cell, space_pts=2)


def generate_trade_document(
    region,
    country_or_group,
    start_year,
    end_year,
    digit=4,
    category=None,
    text_size=7,
    table_size=25,
    country_table_size=15,
    month_range_raw="",
    exclude_raw=""
):
    load_dotenv()
    exclude_tn_veds = [item.strip() for item in exclude_raw.split(",") if item.strip()]
    parts = [int(m.strip()) for m in month_range_raw.split(",") if m.strip()]
    if len(parts) == 2:
        month_range = list(range(parts[0], parts[1] + 1))
    elif len(parts) == 1:
        month_range = [parts[0]]
    else:
        month_range = []

    conn = psycopg2.connect(
        host=os.getenv("DB_HOST"),
        port=os.getenv("DB_PORT"),
        database=os.getenv("DB_NAME"),
        user=os.getenv("DB_USER"),
        password=os.getenv("DB_PASS")
    )

    tradeDataFetcher = TradeDataFetcher(conn)
    if not tradeDataFetcher.is_data_exists(country_or_group, region, end_year, month_range):
        return "Данных нет", "Данных нет", "Данных нет"

    tradeDataPreparer = TradeDataPreparer(conn, region, country_or_group, start_year, end_year, digit, category, text_size, table_size, country_table_size, exclude_tn_veds, month_range)
    
    data_for_doc = tradeDataPreparer.prepare()
    
    tradeDocumentGenerator = TradeDocumentGenerator(data_for_doc)
    
    doc, filename, short_filename = tradeDocumentGenerator.generate()
    if filename == "Данных нет":
        return "Данных нет", "Данных нет", "Данных нет"
    return doc, filename, short_filename
